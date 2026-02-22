import zipfile
import xml.etree.ElementTree as ET
import docx

def read_docx(path, out_file):
    out_file.write(f"--- DOCX: {path} ---\n")
    doc = docx.Document(path)
    for p in doc.paragraphs:
        if p.text.strip():
            out_file.write(p.text + "\n")
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.replace("\n", " ").strip())
            if any(row_data):
                out_file.write(" | ".join(row_data) + "\n")

def read_xlsx(path, out_file):
    out_file.write(f"\n--- XLSX: {path} ---\n")
    try:
        with zipfile.ZipFile(path, 'r') as z:
            for filename in z.namelist():
                if filename.endswith('.xml'):
                    try:
                        xml_content = z.read(filename)
                        root = ET.fromstring(xml_content)
                        for elem in root.iter():
                            if elem.text and elem.text.strip():
                                out_file.write(elem.text.strip() + "\n")
                    except Exception:
                        pass
    except Exception as e:
        out_file.write(f"Failed to read xlsx: {e}\n")

if __name__ == "__main__":
    with open(r"e:\محرر\extracted_text.txt", "w", encoding="utf-8") as f:
        read_docx(r"e:\محرر\docs\competitive-analysis-avan-titre.docx", f)
        read_xlsx(r"e:\محرر\docs\competitive-analysis-avan-titre.xlsx", f)
